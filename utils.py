import tiktoken
import pdfplumber
import hashlib
import re
from typing import List, Dict, Tuple, Optional
from constants import (
    PLACEHOLDER_PATTERN, 
    INPUT_TYPE_KEYWORDS, 
    PLACEHOLDER_DEFAULTS,
    SUPPORTED_FILE_TYPES
)

def count_tokens(text: str) -> int:
    """Count tokens in text using tiktoken."""
    try:
        encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
        return len(encoding.encode(text))
    except Exception as e:
        # Return 0 on error - let the calling function handle the error display
        return 0

def extract_text_from_pdf(pdf_file) -> str:
    """Extract text from uploaded PDF file."""
    try:
        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
    except Exception as e:
        raise Exception(f"Error reading PDF: {e}")

def calculate_content_hash(resume_text: str, job_description: str) -> str:
    """Calculate hash of resume and job description for cache invalidation."""
    content = f"{resume_text}|{job_description}"
    return hashlib.md5(content.encode()).hexdigest()

def extract_placeholders(text: str) -> List[str]:
    """Extract placeholders from text using regex."""
    matches = re.findall(PLACEHOLDER_PATTERN, text)
    return list(set(matches))  # Remove duplicates

def determine_input_type(placeholder: str) -> str:
    """Determine the appropriate input type for a placeholder."""
    placeholder_lower = placeholder.lower()
    
    for input_type, keywords in INPUT_TYPE_KEYWORDS.items():
        if any(keyword in placeholder_lower for keyword in keywords):
            return input_type
    
    return 'text'

def get_placeholder_default(placeholder: str) -> str:
    """Get default value for a placeholder based on its type."""
    input_type = determine_input_type(placeholder)
    return PLACEHOLDER_DEFAULTS.get(input_type, PLACEHOLDER_DEFAULTS['generic'])

def replace_placeholders(text: str, placeholder_values: Dict[str, str]) -> str:
    """Replace placeholders in text with user-provided values."""
    for placeholder, value in placeholder_values.items():
        pattern = f'\\[{re.escape(placeholder)}\\]'
        text = re.sub(pattern, value, text)
    return text

def validate_file_type(filename: str) -> bool:
    """Validate if uploaded file is of supported type."""
    if not filename:
        return False
    
    file_extension = filename.split('.')[-1].lower()
    return file_extension in SUPPORTED_FILE_TYPES

def format_token_count(count: int) -> str:
    """Format token count for display."""
    if count >= 1000:
        return f"{count/1000:.1f}K"
    return str(count)

def get_input_widget_type(placeholder: str) -> str:
    """Get the appropriate Streamlit input widget type for a placeholder."""
    input_type = determine_input_type(placeholder)
    
    if input_type == 'email':
        return 'email'
    elif input_type == 'phone':
        return 'tel'
    elif input_type == 'date':
        return 'date'
    else:
        return 'text'

def sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe file operations."""
    # Remove or replace unsafe characters
    unsafe_chars = '<>:"/\\|?*'
    for char in unsafe_chars:
        filename = filename.replace(char, '_')
    return filename

def truncate_text(text: str, max_length: int = 100) -> str:
    """Truncate text to specified length with ellipsis."""
    if len(text) <= max_length:
        return text
    return text[:max_length-3] + "..."

def validate_required_fields(filled_values: Dict[str, str], required_placeholders: List[str]) -> Tuple[bool, List[str]]:
    """Validate that all required fields are filled."""
    empty_fields = []
    for placeholder in required_placeholders:
        if not filled_values.get(placeholder, '').strip():
            empty_fields.append(placeholder)
    
    return len(empty_fields) == 0, empty_fields

def convert_cover_letter_to_docx(cover_letter: str, personal_info: Dict[str, str]) -> bytes:
    """Convert cover letter and personal info to Word document format."""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Personalized Cover Letter', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add personal information section
    doc.add_heading('Personal Information', level=1)
    for placeholder, value in personal_info.items():
        p = doc.add_paragraph()
        p.add_run(f'{placeholder}: ').bold = True
        p.add_run(value)
    
    # Add separator
    doc.add_paragraph()
    
    # Add cover letter content
    doc.add_heading('Cover Letter Content', level=1)
    
    # Split cover letter into paragraphs and add them
    paragraphs = cover_letter.split('\n\n')
    for paragraph in paragraphs:
        if paragraph.strip():  # Only add non-empty paragraphs
            # Clean up the paragraph (remove extra whitespace)
            clean_paragraph = paragraph.strip()
            if clean_paragraph:
                doc.add_paragraph(clean_paragraph)
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue() 